# -*- coding: utf-8 -*-
# @Time    : 2023/1/11 18:05
# @Author  : HongFei Wang

import tkinter as tk
from tkinter import messagebox
from tkinter.filedialog import *

import docx
import xlrd
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.shared import RGBColor
from tqdm import tqdm


def write(name, pos, idCard, newfile):
	# 增加一级标题, 仿宋居中
	Head = newfile.add_heading("", level=1)  # level设置N级标题
	Head.alignment = WD_ALIGN_PARAGRAPH.CENTER
	run = Head.add_run('关于身份证、户口簿复印件收集情况说明')
	run.font.name = u'宋体'
	run.font.size = Pt(22)
	run.font.color.rgb = RGBColor(0, 0, 0)
	run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
	Head.paragraph_format.line_spacing = 2.1  # 行间距，1.5倍行距
	Head.paragraph_format.space_before = Pt(17)  # 段前30磅
	Head.paragraph_format.space_after = Pt(16.5)  # 段后15磅
	# 添加文字
	text_titel = "按大足区农村承包地确权项目要求，本项目开展至今由于"
	# pos = "大足区智凤街道八里村一组"
	# name = "龚国才"
	# idCard = "500225198912301918"
	
	par1 = newfile.add_paragraph()
	
	par1.paragraph_format.line_spacing = 1.5  # 行间距，1.0倍行距
	par1.paragraph_format.first_line_indent = 15 * 2
	run1 = par1.add_run(text_titel)
	par1.alignment = WD_ALIGN_PARAGRAPH.LEFT
	run1.font.name = u'仿宋'
	run1.font.size = Pt(15)
	run1.font.color.rgb = RGBColor(0, 0, 0)
	run1._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
	par1.paragraph_format.first_line_indent = run1.font.size * 2
	
	run2 = par1.add_run(pos)
	par1.alignment = WD_ALIGN_PARAGRAPH.LEFT
	run2.font.name = u'仿宋'
	run2.font.size = Pt(15)
	run2.font.color.rgb = RGBColor(0, 0, 0)
	run2._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
	run2.underline = True
	
	run2 = par1.add_run("成员")
	par1.alignment = WD_ALIGN_PARAGRAPH.LEFT
	run2.font.name = u'仿宋'
	run2.font.size = Pt(15)
	run2.font.color.rgb = RGBColor(0, 0, 0)
	run2._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
	run2.underline = False
	
	run2 = par1.add_run(f'{name}({idCard})')
	par1.alignment = WD_ALIGN_PARAGRAPH.LEFT
	run2.font.name = u'仿宋'
	run2.font.size = Pt(15)
	run2.font.color.rgb = RGBColor(0, 0, 0)
	run2._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
	run2.underline = True
	
	run2 = par1.add_run('因长期不在家等原因，经村委会多方联络，暂时收集不到户主身份证复印件、家庭成员户口簿复印件。')
	par1.alignment = WD_ALIGN_PARAGRAPH.LEFT
	run2.font.name = u'仿宋'
	run2.font.size = Pt(15)
	run2.font.color.rgb = RGBColor(0, 0, 0)
	run2._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
	run2.underline = False
	
	par2 = newfile.add_paragraph()
	par2.paragraph_format.line_spacing = 1.0  # 行间距，1.0倍行距
	par2.paragraph_format.first_line_indent = 15 * 2
	run2 = par2.add_run('特此说明!')
	par2.alignment = WD_ALIGN_PARAGRAPH.LEFT
	run2.font.name = u'仿宋'
	run2.font.size = Pt(15)
	run2.font.color.rgb = RGBColor(0, 0, 0)
	run2._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
	run2.underline = False
	par2.paragraph_format.first_line_indent = run1.font.size * 2
	
	# 添加空段落
	
	par2 = newfile.add_paragraph()
	par2.paragraph_format.line_spacing = 1.0  # 行间距，1.0倍行距
	run2 = par2.add_run('    ')
	run2.font.name = u'仿宋'
	run2.font.size = Pt(15)
	run2.font.color.rgb = RGBColor(0, 0, 0)
	run2._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
	run2.underline = False
	# 添加落款
	par2 = newfile.add_paragraph()
	par2.paragraph_format.line_spacing = 1.0  # 行间距，1.0倍行距
	run2 = par2.add_run(pos)
	par2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
	run2.font.name = u'仿宋'
	run2.font.size = Pt(15)
	run2.font.color.rgb = RGBColor(0, 0, 0)
	run2._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
	par2.paragraph_format.first_line_indent = run2.font.size * 15
	run2.underline = False
	# 空段落
	par2 = newfile.add_paragraph()
	par2.paragraph_format.line_spacing = 1.0  # 行间距，1.0倍行距
	run2 = par2.add_run('      ')
	run2.font.name = u'仿宋'
	run2.font.size = Pt(15)
	run2.font.color.rgb = RGBColor(0, 0, 0)
	run2._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
	run2.underline = False
	
	# 盖章段落
	par2 = newfile.add_paragraph()
	par2.paragraph_format.line_spacing = 1.0  # 行间距，1.0倍行距
	run2 = par2.add_run("(盖章)")
	par2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
	run2.font.name = u'仿宋'
	run2.font.size = Pt(15)
	run2.font.color.rgb = RGBColor(0, 0, 0)
	run2._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
	run2.underline = False
	par2.paragraph_format.first_line_indent = run2.font.size * 19
	
	# 添加时间
	par2 = newfile.add_paragraph()
	par2.paragraph_format.line_spacing = 1.0  # 行间距，1.0倍行距
	run2 = par2.add_run("2020年12月20日")
	par2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
	run2.font.name = u'仿宋'
	run2.font.size = Pt(15)
	run2.font.color.rgb = RGBColor(0, 0, 0)
	run2._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
	par2.paragraph_format.first_line_indent = run2.font.size * 16
	run2.underline = False
	
	# 开始跳到下一页
	
	for i in range(11):
		run2 = par2.add_run("")
		par2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
		run2.font.name = u'仿宋'
		run2.font.size = Pt(15)
		run2.font.color.rgb = RGBColor(0, 0, 0)
		run2._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
		par2 = newfile.add_paragraph()


def main(**kwargs):
	assert kwargs["excelFile"], messagebox.showerror('错误', 'excel路径出错,得到的是{}'.format(kwargs["excelFile"]))
	assert kwargs["save_path"], messagebox.showerror('错误', '保存路径出错得到的是{}'.format(kwargs["save_path"]))
	try:
		excels = xlrd.open_workbook(kwargs["excelFile"])
		sheets = excels.sheet_by_index(0)
		nrow = sheets.nrows
	except:
		print("读取excel错误")
		nrow = 1
	newfile = docx.Document()
	for row in tqdm(range(1, nrow)):
		values = sheets.row_values(row)
		pos, name, idCard = values
		write(name, pos, idCard, newfile)
	
	newfile.save(kwargs["save_path"])
	messagebox.showinfo('完成', 'word 已经保存到{}'.format(kwargs["save_path"]))


def file(s):
	pa = askopenfilename(filetypes=[('excels', '*.xls')])
	print(pa)
	s.set(pa)


def save(s):
	pa = asksaveasfilename(defaultextension='.docx')
	s.set(pa)


def gui():
	root = tk.Tk()
	root.geometry("300x200+600+300")
	filePath = tk.StringVar()
	savePath = tk.StringVar()
	tk.Label(root, text="     ").grid(row=0, column=0)
	tk.Entry(root, textvariable=filePath).grid(row=0, column=1)
	tk.Button(root, text="选择excel", command=lambda: file(filePath)).grid(row=0, column=2)
	tk.Label(root, text="     ").grid(row=0, column=0)
	tk.Entry(root, textvariable=savePath).grid(row=1, column=1)
	tk.Button(root, text="保存路径", command=lambda: save(savePath)).grid(row=1, column=2)
	tk.Button(root, text=" 运行 ", command=lambda: main(excelFile=filePath.get(), save_path=savePath.get())).grid(row=3, column=1)
	root.mainloop()


if __name__ == '__main__':
	gui()