# -*- coding: utf-8 -*-
# @Time    : 2023/2/7 10:52
# @Author  : HongFei Wang

import tqdm
from docx import Document
from docxcompose.composer import Composer


def delete_paragraph(paragraph):
	p = paragraph._element
	p.getparent().remove(p)
	paragraph._p = paragraph._element = None


# 定义合并文档的函数
def merge_doc(source_file_path_list, target_file_path):
	"""
	:param source_file_path_list: 源文件路径列表
	:param target_file_path: 目标文件路径
	"""
	# 填充分页符号文档
	page_break_doc = Document()
	page_break_doc.add_section()
	# 定义新文档
	target_doc = Document(source_file_path_list[0])
	target_composer = Composer(target_doc)
	for i in tqdm.tqdm(range(len(source_file_path_list))):
		# 跳过第一个作为模板的文件
		if i == 0: continue
		target_composer.append(page_break_doc)  # 添加空白
		# 拼接文档内容
		f = source_file_path_list[i]
		f = Document(f)
		target_composer.append(f)
	# 保存目标文档
	target_composer.save(target_file_path)